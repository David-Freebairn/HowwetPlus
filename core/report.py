"""
core/report.py
===============
Builds a downloadable Word (.docx) summary of a single Howwet?+ run:
inputs used, the fallow and in-crop water balance tables (with historical
percentile ranks where available), and the soil water chart exactly as
shown in the app.

Uses python-docx directly since this runs inside the deployed Streamlit
app itself (on demand, per user, per run) rather than being a one-off
document generated at authoring time.
"""

import io
from datetime import datetime


_COMPONENT_ROWS = [
    ("Rainfall", "rain", "rain"),
    ("Runoff", "runoff", "runoff"),
    ("Soil evaporation", "soil_evap", "soil_evap"),
    ("Transpiration", "transp", "transp"),
    ("Deep drainage", "drainage", "drainage"),
    ("Change in soil water", "dsw", "dsw"),
]


def _phase_totals(sub_df):
    if sub_df is None or sub_df.empty:
        return None
    rain_t = float(sub_df["rain"].sum())
    return {
        "rain": rain_t,
        "runoff": float(sub_df["runoff"].sum()),
        "soil_evap": float(sub_df["soil_evap"].sum()),
        "transp": float(sub_df["transp"].sum()),
        "drainage": float(sub_df["drainage"].sum()),
        "dsw": float(sub_df["sw_total"].iloc[-1] - sub_df["sw_total"].iloc[0]),
        "_rain_total": rain_t,
    }


def _percentile_rank(value, historical_values):
    if not historical_values:
        return None
    import numpy as np
    arr = np.asarray(historical_values, dtype=float)
    return float(100.0 * (arr <= value).mean())


def _add_water_balance_table(doc, title, sub_df, hist_phase, extra_row=None):
    doc.add_heading(title, level=2)
    totals = _phase_totals(sub_df)
    if totals is None:
        doc.add_paragraph("No days in this run for this phase.")
        return

    has_hist = bool(hist_phase)
    ncols = 4 if has_hist else 3
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "mm"
    hdr[2].text = "% of rainfall"
    if has_hist:
        hdr[3].text = "Historical %ile"

    rain_total = totals["_rain_total"]
    for label, key, hist_key in _COMPONENT_ROWS:
        val = totals[key]
        pct_rain = 100.0 if key == "rain" else (val / rain_total * 100.0 if rain_total > 0 else 0.0)
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{val:.0f}"
        row[2].text = f"{pct_rain:.0f}"
        if has_hist:
            p = _percentile_rank(val, hist_phase.get(hist_key)) if hist_key else None
            row[3].text = f"{p:.0f}th" if p is not None else "\u2014"

    if extra_row is not None:
        label, val, pctile_str = extra_row
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{val:.0f}"
        row[2].text = "\u2014"
        if has_hist:
            row[3].text = pctile_str if pctile_str is not None else "\u2014"


def build_report_docx(inputs: dict, df, hist_pct: dict, plume: dict, chart_png: bytes,
                       icon_path=None, fallow_n_actual=None, fallow_n_hist=None,
                       n_chart_png=None, budget=None) -> bytes:
    """
    Parameters
    ----------
    inputs    : dict of input labels -> display strings (station, soil,
                crop template, fallow start, plant date, maturity date,
                simulated-to date, starting soil water %, residue cover %,
                crop factor)
    df        : the daily results DataFrame from run_fallow_to_crop
    hist_pct  : historical_phase_percentiles() output, or None
    plume     : pasw_plume_from_replays() output, or None (used for the
                "N replayed seasons" note)
    chart_png : PNG bytes of the soil water chart, exactly as shown in-app
    icon_path : optional path to a logo image, placed above the title
    fallow_n_hist : optional list of this season's fallow N mineralisation
                total (kg/ha) — one value per historical replayed year —
                from core.nitrogen.fallow_n_historical_values(). Adds a
                "Gain in nitrate" row to the Fallow table with the current
                season's actual value ranked against this distribution.
    n_chart_png : optional PNG bytes of the Nitrate mineralisation chart,
                exactly as shown in-app. Adds a "Nitrate mineralisation"
                section with this image when given.
    budget    : optional core.yield_n.yield_n_budget() output (dict keyed
                20/50/80). Adds a "Yield & nitrogen calculator table"
                section when given.

    Returns
    -------
    bytes of the generated .docx file
    """
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    if icon_path:
        try:
            icon_para = doc.add_paragraph()
            icon_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            icon_para.add_run().add_picture(str(icon_path), width=Inches(1.0))
        except Exception:
            pass  # missing/unreadable icon shouldn't block the report

    doc.add_heading("Howwet? + \u2014 Soil Water Report", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated {datetime.now().strftime('%d %b %Y %H:%M')}").italic = True

    doc.add_heading("Inputs", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in inputs.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    fallow_df = df[df["phase"] == "fallow"]
    crop_df   = df[df["phase"] == "crop"]
    hist_fallow = hist_pct["fallow"] if hist_pct else None
    hist_crop   = hist_pct["crop"] if hist_pct else None

    fallow_n_extra = None
    if fallow_n_actual is not None:
        pctile_str = None
        if fallow_n_hist:
            p = _percentile_rank(fallow_n_actual, fallow_n_hist)
            pctile_str = f"{p:.0f}th" if p is not None else None
        fallow_n_extra = ("Gain in nitrate (kg/ha)", fallow_n_actual, pctile_str)

    _add_water_balance_table(doc, "Water balance \u2014 Fallow", fallow_df, hist_fallow, extra_row=fallow_n_extra)
    _add_water_balance_table(doc, "Water balance \u2014 In-crop", crop_df, hist_crop)

    if plume:
        note = doc.add_paragraph()
        note.add_run(
            f"Historical percentile ranks are against {plume['n_years']} replayed seasons "
            f"since 1995 at this site, using the same fallow/plant/maturity month-day pattern "
            f"applied to each year's actual climate record."
        ).italic = True

    doc.add_heading("Soil water", level=2)
    if chart_png:
        doc.add_picture(io.BytesIO(chart_png), width=Inches(6.5))

    if n_chart_png:
        doc.add_heading("Nitrate mineralisation", level=2)
        doc.add_picture(io.BytesIO(n_chart_png), width=Inches(6.5))

    if budget:
        doc.add_heading("Yield & nitrogen calculator table", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = ""
        hdr[1].text = "20%ile"
        hdr[2].text = "50%ile"
        hdr[3].text = "80%ile"

        def _row(label, vals):
            row = table.add_row().cells
            row[0].text = label
            for i, v in enumerate(vals):
                row[i + 1].text = v

        _row("Estimated water-limited yield (t/ha)",
             [f"{budget[p]['yield_t_ha']:.1f}" for p in (20, 50, 80)])
        _row("Nitrogen required (kg N/ha)",
             [f"{budget[p]['n_required']:.0f}" for p in (20, 50, 80)])
        _row("Nitrogen supply (kg N/ha)",
             [f"{budget[p]['n_supply']:.0f}" for p in (20, 50, 80)])
        _row("Nitrogen deficit/surplus (kg N/ha)",
             [f"{budget[p]['n_balance']:+.0f}" for p in (20, 50, 80)])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
